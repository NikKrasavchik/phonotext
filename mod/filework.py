# -*- coding: utf-8 -*-
'''
Работа с файлами
'''


def txt_out(name: str, analyzed_text: dict):
    '''
    Вывод результата в .txt файл
    '''
    with open('archive/' + name + '.txt', 'wt', encoding='utf-8') as file:
        file.write('Базовый текст\n\n')
        skip = 0
        for line in analyzed_text['lines']:
            text_line = ''
            for span in line:
                text_line = text_line + span[0]
            if text_line == '':
                if skip:
                    skip = 0
                    file.write(str(text_line) + '\n')
                else:
                    skip = 1
            else:
                skip = 0
                file.write(str(text_line) + '\n')

        file.write('\n\n\n')
        file.write('Фоносиллабы\n\n')
        syll_text = ''
        for syll in analyzed_text['syll']:
            if '/' in syll:
                syll_text = syll_text + '\n'
            else:
                syll_text = syll_text + syll[0] + ' '
        file.write(str(syll_text) + '\n')

        file.write('\n\n\n')
        file.write('Комбинации\n\n')
        current = 0
        comb_line = ''
        comb_iter = 0
        for comb in analyzed_text['comb']:
            if comb[0] != current:
                file.write(str(comb_line) + '\n')
                current = comb[0]
                comb_iter = comb_iter + 1
                comb_line = str(comb_iter) + ') ' + comb[1] + ', '
            else:
                comb_line = comb_line + comb[1] + ', '
        file.write(str(comb_line) + '\n')

        file.write('\n\n\n')
        file.write('Вибрации\n\n')
        for vibr in analyzed_text['vibr']:
            file.write(str(vibr[1]) + ' в позиции ' + str(vibr[0]) + '\n')

        file.write('\n\n\n')
        file.write('Повторы\n\n')
        for i, rep in enumerate(analyzed_text['rep']):
            file.write(str(i) + '). ')
            for single in rep:
                file.write(single[1] + ', ')


def doc_out(name: str, analyzed_text: dict):
    '''
    Вывод результата в .doc файл
    '''
    import docx

    document = docx.Document()

    document.add_heading('Базовый текст', 0)
    skip = 0
    for line in analyzed_text['lines']:
        text_line = ''
        for span in line:
            text_line = text_line + span[0]
        if text_line == '':
            if skip:
                skip = 0
                document.add_paragraph(str(text_line))
            else:
                skip = 1
        else:
            skip = 0
            document.add_paragraph(str(text_line))

    document.add_page_break()
    document.add_heading('Фоносиллабы', 0)
    syll_text = ''
    for syll in analyzed_text['syll']:
        if '/' in syll:
            syll_text = syll_text + '\n'
        else:
            syll_text = syll_text + syll[0] + ' '
    document.add_paragraph(str(syll_text))

    document.add_page_break()
    document.add_heading('Комбинации', 0)
    current = 0
    comb_line = ''
    comb_iter = 0
    for comb in analyzed_text['comb']:
        if comb[0] != current:
            document.add_paragraph(str(comb_line))
            current = comb[0]
            comb_iter = comb_iter + 1
            comb_line = str(comb_iter) + ') ' + comb[1] + ', '
        else:
            comb_line = comb_line + comb[1] + ', '
    document.add_paragraph(str(comb_line))

    document.add_page_break()
    document.add_heading('Вибрации', 0)
    for vibr in analyzed_text['vibr']:
        document.add_paragraph(str(vibr[1]) + ' в позиции ' + str(vibr[0]))

    document.add_page_break()
    document.add_heading('Повторы', 0)
    for i, rep in enumerate(analyzed_text['rep']):
        temp = str(i) + '). '
        for single in rep:
            temp += single[1] + ', '
        document.add_paragraph(temp)

    document.save('archive/' + name + '.docx')
